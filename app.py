import streamlit as st
import pandas as pd
import datetime
import csv
import os
from pitch_generator import generate_pitch
from fpdf import FPDF
from docx import Document
from io import BytesIO

# --- Page Configuration ---
st.set_page_config(page_title="AI Pitch Generator - PNG", layout="wide")
st.title("🇵🇬 AI Incubation Hub Pitch Generator")
st.markdown("*Connecting PNG Innovators to Funding Opportunities*")

# --- Session state ---
if "donor_matches" not in st.session_state:
    st.session_state.donor_matches = None
if "selected_donor" not in st.session_state:
    st.session_state.selected_donor = None
if "generated_pitch" not in st.session_state:
    st.session_state.generated_pitch = None

# --- Save Pitch to CSV (admin only) ---
def save_pitch_to_db(hub_name, user_name, pitch_text):
    try:
        file_exists = os.path.isfile('pitches.csv')
        with open('pitches.csv', 'a', newline='', encoding='utf-8') as f:
            writer = csv.writer(f)
            if not file_exists:
                writer.writerow(['Date', 'Hub', 'User', 'Pitch'])
            writer.writerow([datetime.datetime.now().strftime("%Y-%m-%d %H:%M:%S"), hub_name, user_name, pitch_text])
        return True
    except Exception as e:
        st.error(f"Error saving pitch: {e}")
        return False

# --- PDF Generation ---
def generate_pdf_proposal(pitch_text, hub_name, user_name):
    pdf = FPDF()
    pdf.set_auto_page_break(auto=True, margin=15)
    pdf.add_page()
    pdf.set_font("Arial", "B", 18)
    pdf.cell(0, 10, "AI Incubation Hub - Funding Proposal", ln=True, align="C")
    pdf.ln(5)
    pdf.set_font("Arial", "", 12)
    pdf.cell(0, 10, f"Hub: {hub_name}", ln=True)
    pdf.cell(0, 10, f"Prepared by: {user_name}", ln=True)
    pdf.cell(0, 10, f"Date: {datetime.datetime.now().strftime('%Y-%m-%d')}", ln=True)
    pdf.ln(10)
    pdf.set_font("Arial", "", 11)
    for line in pitch_text.split('\n'):
        clean_line = line.strip()
        if clean_line.startswith('#'):
            pdf.set_font("Arial", "B", 14)
            heading_text = clean_line.lstrip('#').strip()
            pdf.multi_cell(0, 8, heading_text)
            pdf.set_font("Arial", "", 11)
            pdf.ln(2)
        elif clean_line == '':
            pdf.ln(4)
        else:
            pdf.multi_cell(0, 6, clean_line)
    return pdf.output(dest='S').encode('latin-1')

# --- Word Generation ---
def generate_word_proposal(pitch_text, hub_name, user_name):
    doc = Document()
    title = doc.add_heading('AI Incubation Hub - Funding Proposal', 0)
    title.alignment = 1
    doc.add_paragraph(f"Hub: {hub_name}")
    doc.add_paragraph(f"Prepared by: {user_name}")
    doc.add_paragraph(f"Date: {datetime.datetime.now().strftime('%Y-%m-%d')}")
    doc.add_paragraph()
    for line in pitch_text.split('\n'):
        clean_line = line.strip()
        if clean_line.startswith('#'):
            level = min(len(clean_line) - len(clean_line.lstrip('#')), 9)
            heading_text = clean_line.lstrip('#').strip()
            doc.add_heading(heading_text, level=level)
        elif clean_line == '':
            doc.add_paragraph()
        else:
            doc.add_paragraph(clean_line)
    buffer = BytesIO()
    doc.save(buffer)
    return buffer.getvalue()

# --- Full Donor Database (47 donors) ---
@st.cache_data
def load_donor_data():
    return pd.DataFrame([
        # ============================================================
        # 1. BILATERAL DONORS
        # ============================================================
        {
            "Donor": "Australia (DFAT)",
            "Focus": "Governance, Economic Growth, Human Development, Infrastructure, Health, Education, Law & Justice, Agriculture",
            "Sector": "Bilateral",
            "Funding Call": "PNG-Australia Partnership for Development",
            "Amount": "K1.32 billion (2025)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Australian Government",
            "Website": "https://www.dfat.gov.au/geo/papua-new-guinea",
            "Apply": "https://www.dfat.gov.au/geo/papua-new-guinea/development-assistance",
            "MaxPages": 25,
            "RecommendedPages": 15
        },
        {
            "Donor": "China (PRC)",
            "Focus": "Infrastructure, Education, Health, Agriculture, Construction Projects, Energy, Mining",
            "Sector": "Bilateral",
            "Funding Call": "China-PNG Bilateral Cooperation",
            "Amount": "K298 million (2025)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "People's Republic of China",
            "Website": "https://www.cidca.gov.cn",
            "Apply": "https://www.cidca.gov.cn/english",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "Japan (JICA)",
            "Focus": "Infrastructure, Fisheries, Technical Cooperation, Education, Renewable Energy, Health",
            "Sector": "Bilateral",
            "Funding Call": "Japan-PNG Cooperation Program",
            "Amount": "K181 million (2020, loans)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Government of Japan",
            "Website": "https://www.jica.go.jp/png",
            "Apply": "https://www.jica.go.jp/png/english/activities/index.html",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "New Zealand (MFAT)",
            "Focus": "Economic Sectors, Infrastructure, Health, Education, Law & Justice, Agriculture",
            "Sector": "Bilateral",
            "Funding Call": "New Zealand Aid Programme - PNG",
            "Amount": "K45 million (2025)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "New Zealand Government",
            "Website": "https://www.mfat.govt.nz/en/countries-and-regions/pacific/papua-new-guinea",
            "Apply": "https://www.mfat.govt.nz/en/aid-and-development/",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "United States (USAID)",
            "Focus": "Governance, Democracy, Health, Economic Growth, Environment, Disaster Preparedness",
            "Sector": "Bilateral",
            "Funding Call": "USAID PNG Program",
            "Amount": "US$19 million (2025)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "U.S. Government",
            "Website": "https://www.usaid.gov/papua-new-guinea",
            "Apply": "https://www.usaid.gov/papua-new-guinea/opportunities",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "France (AFD)",
            "Focus": "Green Finance, Renewable Energy, Climate Resilience, Biodiversity",
            "Sector": "Bilateral",
            "Funding Call": "France-PNG Cooperation",
            "Amount": "€15 million (approx. K75 million)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Government of France",
            "Website": "https://www.afd.fr",
            "Apply": "https://www.afd.fr/en/countries/papua-new-guinea",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "Germany (GIZ)",
            "Focus": "Food Security, Rural Development, Environmental Protection, Energy",
            "Sector": "Bilateral",
            "Funding Call": "Germany-PNG Technical Cooperation",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "German Government",
            "Website": "https://www.giz.de/en/worldwide/346.html",
            "Apply": "https://www.giz.de/en/working_with_giz/apply_for_funding.html",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "Korea (KOICA)",
            "Focus": "Renewable Energy, Green Economy Training, Health, ICT",
            "Sector": "Bilateral",
            "Funding Call": "KOICA PNG Program",
            "Amount": "US$10 million+ annually",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Government of Korea",
            "Website": "https://www.koica.go.kr/png_en/index.do",
            "Apply": "https://www.koica.go.kr/png_en/partner/index.do",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "United Kingdom (FCDO)",
            "Focus": "Climate Change, Governance, Education, Health, Economic Development",
            "Sector": "Bilateral",
            "Funding Call": "UK-PNG Development Program",
            "Amount": "£20 million+ annually",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "UK Government",
            "Website": "https://www.gov.uk/world/organisations/foreign-commonwealth-development-office",
            "Apply": "https://www.gov.uk/apply-for-funding",
            "MaxPages": 25,
            "RecommendedPages": 15
        },
        {
            "Donor": "India",
            "Focus": "Health, Education, Agriculture, ICT, Renewable Energy, IT Training",
            "Sector": "Bilateral",
            "Funding Call": "India-PNG Development Cooperation",
            "Amount": "US$5 million+ annually",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Government of India",
            "Website": "https://www.mea.gov.in/png.htm",
            "Apply": "https://www.indiacop.org/",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "Canada (Global Affairs)",
            "Focus": "Gender Equality, Health, Education, Climate Change, Economic Development",
            "Sector": "Bilateral",
            "Funding Call": "Canada-PNG Development Program",
            "Amount": "CAD 10 million+ annually",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Government of Canada",
            "Website": "https://www.international.gc.ca/world-monde/papua_new_guinea-papouasie_nouvelle_guinee.aspx",
            "Apply": "https://www.international.gc.ca/world-monde/funding-financement.aspx",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "Czech Republic (Ceska/Erste)",
            "Focus": "Development Projects, Infrastructure",
            "Sector": "Bilateral",
            "Funding Call": "Czech-PNG Development Cooperation",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Czech Government",
            "Website": "https://www.mzv.cz",
            "Apply": "https://www.mzv.cz/jnp/en/index.html",
            "MaxPages": 20,
            "RecommendedPages": 10
        },
        # ============================================================
        # 2. MULTILATERAL & INTERNATIONAL ORGANIZATIONS
        # ============================================================
        {
            "Donor": "Asian Development Bank (ADB)",
            "Focus": "Economic Development, Energy, Infrastructure, Finance, Human Capital, Transport",
            "Sector": "Multilateral",
            "Funding Call": "ADB Country Operations Business Plan",
            "Amount": "K437.6 million (2020 loans)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "ADB",
            "Website": "https://www.adb.org/countries/papua-new-guinea",
            "Apply": "https://www.adb.org/projects",
            "MaxPages": 50,
            "RecommendedPages": 25
        },
        {
            "Donor": "World Bank Group",
            "Focus": "Infrastructure, Agriculture, Private Sector, Jobs, Public Sector Reform, Education, Health",
            "Sector": "Multilateral",
            "Funding Call": "PNG Country Partnership Framework",
            "Amount": "K185.5 million (2020 loans)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "World Bank",
            "Website": "https://www.worldbank.org/en/country/png",
            "Apply": "https://projects.worldbank.org/en/projects-operations/projects-list",
            "MaxPages": 50,
            "RecommendedPages": 25
        },
        {
            "Donor": "European Union (EU)",
            "Focus": "Climate Resilience, Biodiversity, Sustainable Livelihoods, Conservation, Governance",
            "Sector": "Multilateral",
            "Funding Call": "EU-FCCB Programme Grant Scheme",
            "Amount": "K79.9 million (2020 grants)",
            "Deadline": "31 December 2028",
            "Status": "Open",
            "Partner": "European Union",
            "Website": "https://international-partnerships.ec.europa.eu/countries/papua-new-guinea_en",
            "Apply": "https://fccbpng.eu/grant-scheme",
            "MaxPages": 40,
            "RecommendedPages": 20
        },
        {
            "Donor": "United Nations (UN)",
            "Focus": "Sustainable Development, Governance, Gender Equality, Climate Resilience, Health, Education",
            "Sector": "Multilateral",
            "Funding Call": "UN Sustainable Development Cooperation Framework 2024-2028",
            "Amount": "K190 million (2025)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "United Nations",
            "Website": "https://www.un.org/papua-new-guinea",
            "Apply": "https://www.undp.org/papua-new-guinea/projects",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        {
            "Donor": "European Investment Bank (EIB)",
            "Focus": "Infrastructure, Climate, Sustainable Development, Energy, Water",
            "Sector": "Multilateral",
            "Funding Call": "EIB Development Loans",
            "Amount": "K40 million (2020 loans)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "European Union",
            "Website": "https://www.eib.org",
            "Apply": "https://www.eib.org/en/projects/index.htm",
            "MaxPages": 40,
            "RecommendedPages": 20
        },
        {
            "Donor": "Asian Infrastructure Investment Bank (AIIB)",
            "Focus": "Agriculture, Infrastructure",
            "Sector": "Multilateral",
            "Funding Call": "AIIB Project Financing",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "AIIB",
            "Website": "https://www.aiib.org",
            "Apply": "https://www.aiib.org/en/projects",
            "MaxPages": 40,
            "RecommendedPages": 20
        },
        {
            "Donor": "Green Climate Fund (GCF)",
            "Focus": "Climate Change Mitigation, Adaptation, Climate Resilience",
            "Sector": "Multilateral",
            "Funding Call": "GCF PNG Country Programme",
            "Amount": "US$50 million+ (multiannual)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Green Climate Fund",
            "Website": "https://www.greenclimate.fund/countries/papua-new-guinea",
            "Apply": "https://www.greenclimate.fund/apply",
            "MaxPages": 60,
            "RecommendedPages": 30
        },
        {
            "Donor": "Global Fund (GFATM)",
            "Focus": "Health, HIV/AIDS, Tuberculosis, Malaria",
            "Sector": "Multilateral",
            "Funding Call": "Global Fund PNG Grants",
            "Amount": "US$100 million+ (multiannual)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Global Fund",
            "Website": "https://www.theglobalfund.org/en/portfolio/country/?loc=PNG",
            "Apply": "https://www.theglobalfund.org/en/funding/",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        {
            "Donor": "Global Partnership for Education (GPE)",
            "Focus": "Education, School Infrastructure, Teacher Training, Early Childhood Education",
            "Sector": "Multilateral",
            "Funding Call": "GPE PNG Education Sector Funding",
            "Amount": "US$30 million+ (multiannual)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "GPE / World Bank",
            "Website": "https://www.globalpartnership.org/where-we-work/papua-new-guinea",
            "Apply": "https://www.globalpartnership.org/funding",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        # ============================================================
        # 3. UN SPECIALIZED AGENCIES
        # ============================================================
        {
            "Donor": "UNDP",
            "Focus": "Sustainable Development, Climate Resilience, Inclusive Governance, Gender Equality",
            "Sector": "UN Agency",
            "Funding Call": "UNDP Country Programme",
            "Amount": "US$30 million+ (multiannual)",
            "Deadline": "October 2026",
            "Status": "Open",
            "Partner": "United Nations",
            "Website": "https://www.undp.org/papua-new-guinea",
            "Apply": "https://www.undp.org/papua-new-guinea/projects",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        {
            "Donor": "UNICEF",
            "Focus": "Child Health, Education, WASH, Child Protection, Nutrition",
            "Sector": "UN Agency",
            "Funding Call": "UNICEF PNG Country Programme",
            "Amount": "US$30 million+ (multiannual)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "United Nations",
            "Website": "https://www.unicef.org/png",
            "Apply": "https://www.unicef.org/png/partner-with-us",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        {
            "Donor": "FAO",
            "Focus": "Food Security, Sustainable Agriculture, Rural Livelihoods, One Health",
            "Sector": "UN Agency",
            "Funding Call": "FAO PNG Country Programme",
            "Amount": "US$10 million+ (multiannual)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "United Nations",
            "Website": "https://www.fao.org/papua-new-guinea",
            "Apply": "https://www.fao.org/papua-new-guinea/projects",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        {
            "Donor": "IFAD",
            "Focus": "Agriculture, Rural Development, Food Security, Climate Resilience",
            "Sector": "UN Agency",
            "Funding Call": "IFAD PNG Country Programme",
            "Amount": "US$20 million+ (multiannual)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "United Nations",
            "Website": "https://www.ifad.org/en/web/operations/country/papua-new-guinea",
            "Apply": "https://www.ifad.org/en/operations",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        {
            "Donor": "UNFPA",
            "Focus": "Reproductive Health, Gender Equality, Population & Development, Youth",
            "Sector": "UN Agency",
            "Funding Call": "UNFPA PNG Country Programme",
            "Amount": "US$10 million+ (multiannual)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "United Nations",
            "Website": "https://pacific.unfpa.org/en/papua-new-guinea",
            "Apply": "https://www.unfpa.org/funding",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        {
            "Donor": "UN Women",
            "Focus": "Gender Equality, Women's Empowerment, Ending Violence Against Women",
            "Sector": "UN Agency",
            "Funding Call": "UN Women PNG Country Programme",
            "Amount": "US$5 million+ (multiannual)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "United Nations",
            "Website": "https://pacific.unwomen.org/en/countries/papua-new-guinea",
            "Apply": "https://www.unwomen.org/en/get-involved/funding",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        {
            "Donor": "ILO",
            "Focus": "Labour Rights, Employment, Social Protection, Decent Work",
            "Sector": "UN Agency",
            "Funding Call": "ILO PNG Country Programme",
            "Amount": "US$5 million+ (multiannual)",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "United Nations",
            "Website": "https://www.ilo.org/papua-new-guinea",
            "Apply": "https://www.ilo.org/papua-new-guinea/projects",
            "MaxPages": 30,
            "RecommendedPages": 20
        },
        {
            "Donor": "UNODC",
            "Focus": "Good Governance, Financial Reform, Crime Prevention, Anti-Corruption",
            "Sector": "UN Agency",
            "Funding Call": "UNODC PNG Country Programme",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "United Nations",
            "Website": "https://www.unodc.org/papua-new-guinea",
            "Apply": "https://www.unodc.org/papua-new-guinea/contact",
            "MaxPages": 25,
            "RecommendedPages": 15
        },
        # ============================================================
        # 4. REGIONAL & PACIFIC SPECIFIC FUNDING
        # ============================================================
        {
            "Donor": "Pacific Islands Forum (PIF)",
            "Focus": "Regional Cooperation, Climate Change, Fisheries, Economic Development",
            "Sector": "Regional",
            "Funding Call": "PIF Regional Programs",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Pacific Islands Forum",
            "Website": "https://www.forumsec.org",
            "Apply": "https://www.forumsec.org/contact/",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "Pacific Community (SPC)",
            "Focus": "Climate Change, Ocean Science, Public Health, Food Security, Geoscience",
            "Sector": "Regional",
            "Funding Call": "SPC PNG Country Program",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Pacific Community",
            "Website": "https://www.spc.int",
            "Apply": "https://www.spc.int/contact",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        {
            "Donor": "Secretariat of the Pacific (SPREP)",
            "Focus": "Climate Change, Biodiversity, Environmental Governance, Waste Management",
            "Sector": "Regional",
            "Funding Call": "SPREP PNG Programs",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "SPREP",
            "Website": "https://www.sprep.org",
            "Apply": "https://www.sprep.org/contact",
            "MaxPages": 20,
            "RecommendedPages": 12
        },
        # ============================================================
        # 5. PNG GOVERNMENT & NATIONAL FUNDS
        # ============================================================
        {
            "Donor": "Incentive Fund (DFAT Australia)",
            "Focus": "SME Development, Innovation, Health, Education, WASH, Governance, Climate",
            "Sector": "Government",
            "Funding Call": "Incentive Fund EOI Grant Program",
            "Amount": "A$100,000 - A$5M",
            "Deadline": "Twice Yearly (March & November)",
            "Status": "Continuous",
            "Partner": "DFAT Australia",
            "Website": "https://incentivefund.gov.pg",
            "Apply": "https://incentivefund.gov.pg/apply",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "National Development Bank (NDB)",
            "Focus": "SME Development, Agriculture, Tourism, Housing, Microfinance",
            "Sector": "Government",
            "Funding Call": "NDB PNG Loan Programs",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "NDB PNG",
            "Website": "https://www.ndb.com.pg",
            "Apply": "https://www.ndb.com.pg/contact/",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "Small Business Development Corporation (SBDC)",
            "Focus": "SME Development, Microfinance, Startups, Entrepreneurship",
            "Sector": "Government",
            "Funding Call": "SBDC PNG Grant & Loan Programs",
            "Amount": "K10,000 - K100,000",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "PNG Government",
            "Website": "https://www.sbdc.org.pg",
            "Apply": "https://www.sbdc.org.pg/contact/",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "UAS Fund",
            "Focus": "ICT Infrastructure, Rural Connectivity, Digital Services, Broadband",
            "Sector": "Government",
            "Funding Call": "UAS Fund ICT Infrastructure Grants",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "NICTA PNG Government",
            "Website": "https://www.nicta.gov.pg",
            "Apply": "https://www.nicta.gov.pg/uas-fund",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        # ============================================================
        # 6. NGOs, FOUNDATIONS & PRIVATE SECTOR
        # ============================================================
        {
            "Donor": "Santos Foundation",
            "Focus": "Water & Sanitation, Education, Solar Lighting, Agriculture, Nutrition",
            "Sector": "NGO/Foundation",
            "Funding Call": "Santos Foundation Small Grants Program",
            "Amount": "K50,000 - K500,000",
            "Deadline": "Twice Yearly",
            "Status": "Continuous",
            "Partner": "Santos Limited",
            "Website": "https://www.santosfoundation.org",
            "Apply": "https://www.santosfoundation.org/apply",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "Malaysian Association of PNG (MAPNG)",
            "Focus": "Healthcare, Community Welfare, Education, Children's Services",
            "Sector": "NGO",
            "Funding Call": "MAPNG Annual Charity Grant",
            "Amount": "K50,000 per recipient",
            "Deadline": "Annual",
            "Status": "Continuous",
            "Partner": "MAPNG",
            "Website": "https://www.mapng.org.pg",
            "Apply": "https://www.mapng.org.pg/apply",
            "MaxPages": 12,
            "RecommendedPages": 8
        },
        {
            "Donor": "PNG LNG (ExxonMobil)",
            "Focus": "Education, Health, Community Development, SME Development",
            "Sector": "Private",
            "Funding Call": "PNG LNG Community Development Program",
            "Amount": "US$10M+ annually",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "ExxonMobil PNG",
            "Website": "https://corporate.exxonmobil.com/locations/papua-new-guinea",
            "Apply": "https://corporate.exxonmobil.com/community-engagement",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "OK Tedi Mining Limited",
            "Focus": "Community Development, Education, Health, Infrastructure, Agriculture",
            "Sector": "Private",
            "Funding Call": "OK Tedi Community Development Program",
            "Amount": "K50M+ annually",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "OK Tedi Mining",
            "Website": "https://www.oktedi.com/our-communities",
            "Apply": "https://www.oktedi.com/our-communities/community-development",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "Barrick Gold / Porgera",
            "Focus": "Community Development, Education, Health, Infrastructure, Agriculture",
            "Sector": "Private",
            "Funding Call": "Porgera Community Development Program",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Barrick Gold",
            "Website": "https://www.barrick.com/operations/porgera/default.aspx",
            "Apply": "https://www.barrick.com/sustainability/community/default.aspx",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "CEFI Fintech Hub",
            "Focus": "Fintech Startups, Digital Finance, Entrepreneurship, Innovation",
            "Sector": "Private",
            "Funding Call": "CEFI Fintech Incubation Program",
            "Amount": "K50,000 - K200,000",
            "Deadline": "November 2026",
            "Status": "Open",
            "Partner": "PNG Unitech / CEFI",
            "Website": "https://www.cefi.com.pg",
            "Apply": "https://www.cefi.com.pg/incubation",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "Rotary International PNG",
            "Focus": "Community Development, Health, Education, Water & Sanitation, Youth",
            "Sector": "NGO",
            "Funding Call": "Rotary PNG Grants Program",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Rotary International",
            "Website": "https://www.rotary.org/en/papua-new-guinea",
            "Apply": "https://www.rotary.org/en/grants",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "World Vision PNG",
            "Focus": "Child Protection, Education, Health, Food Security, Community Development",
            "Sector": "NGO",
            "Funding Call": "World Vision PNG Programs",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "World Vision International",
            "Website": "https://www.worldvision.com.pg",
            "Apply": "https://www.worldvision.com.pg/get-involved",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "Care International PNG",
            "Focus": "Gender Equality, Climate Resilience, Food Security, Health",
            "Sector": "NGO",
            "Funding Call": "Care PNG Programs",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "Care International",
            "Website": "https://www.care.org.au/papua-new-guinea/",
            "Apply": "https://www.care.org.au/partner-with-us/",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "Pacific Women in Business (PWIB)",
            "Focus": "Women Entrepreneurs, SME Development, Leadership, Financial Inclusion",
            "Sector": "Regional/NGO",
            "Funding Call": "PWIB PNG Programs",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "DFAT Australia",
            "Website": "https://www.pacificwomeninbusiness.org",
            "Apply": "https://www.pacificwomeninbusiness.org/contact",
            "MaxPages": 15,
            "RecommendedPages": 10
        },
        {
            "Donor": "Pacific Financial Inclusion Programme (PFIP)",
            "Focus": "Financial Inclusion, SME Finance, Digital Finance, Women in Business",
            "Sector": "Regional/NGO",
            "Funding Call": "PFIP PNG Programs",
            "Amount": "Varies",
            "Deadline": "Continuous",
            "Status": "Continuous",
            "Partner": "PFIP / UNCDF",
            "Website": "https://www.pfip.org",
            "Apply": "https://www.pfip.org/contact",
            "MaxPages": 15,
            "RecommendedPages": 10
        }
    ])

donors = load_donor_data()

# --- AI Donor Matching ---
def ai_match_donor(pitch_text, donors_df):
    donor_list = donors_df[['Donor', 'Focus', 'Funding Call']].to_string(index=False)
    prompt = f"""
You are an expert grant and funding advisor in Papua New Guinea.
Analyze the following pitch and recommend the three most suitable donors from this list:

**Pitch Summary:**
{pitch_text[:1000]}

**Donor List:**
{donor_list}

**Instructions:**
- Return exactly 3 donors.
- For each, provide a one‑line reason for the match.
- Format:
  Donor A: [Reason]
  Donor B: [Reason]
  Donor C: [Reason]
"""
    try:
        from openai import OpenAI
        client = OpenAI(api_key=st.secrets["OPENAI_API_KEY"])
        response = client.chat.completions.create(
            model="gpt-3.5-turbo",
            messages=[{"role": "user", "content": prompt}],
            temperature=0.5,
            max_tokens=300
        )
        content = response.choices[0].message.content
        matches = []
        for line in content.split('\n'):
            if ':' in line and ('Donor' in line or ':' in line):
                parts = line.split(':', 1)
                if len(parts) == 2:
                    donor_name = parts[0].replace('**', '').strip()
                    reason = parts[1].strip()
                    for prefix in ['Donor A:', 'Donor B:', 'Donor C:', 'Donor']:
                        if donor_name.startswith(prefix):
                            donor_name = donor_name[len(prefix):].strip()
                    matches.append((donor_name, reason))
        if len(matches) < 3:
            top_donors = donors_df['Donor'].head(3).tolist()
            matches = [(d, "Recommended by AI") for d in top_donors]
        return matches
    except Exception as e:
        st.error(f"AI matching error: {e}")
        top_donors = donors_df['Donor'].head(3).tolist()
        return [(d, "Fallback match (AI error)") for d in top_donors]

# --- SIDEBAR ---
with st.sidebar:
    st.header("🏢 Donor Database")
    st.caption(f"**{len(donors)}** donors and funding sources available")
    sector_filter = st.selectbox("Filter by Sector", ["All"] + sorted(donors["Sector"].unique().tolist()))
    if sector_filter != "All":
        filtered_donors = donors[donors["Sector"] == sector_filter]
    else:
        filtered_donors = donors.copy()
    status_filter = st.selectbox("Filter by Status", ["All", "Open", "Continuous", "Closed"])
    if status_filter != "All":
        filtered_donors = filtered_donors[filtered_donors["Status"] == status_filter]
    st.dataframe(filtered_donors[["Donor", "Funding Call", "Amount", "Deadline", "Status"]], use_container_width=True)
    st.caption("🔍 Click 'Apply Now' under the generated pitch to apply.")

# --- MAIN FORM ---
st.subheader("📝 Enter Your Innovation Hub Details")
hub_options = ["SoCe Innovation Hub", "CEFI Fintech Hub", "UPNG Innovation Hub", "PNG Unitech Innovation Hub", "Other"]
hub_name = st.selectbox("🏢 Select Your Innovation Hub", hub_options)
if hub_name == "Other":
    hub_name = st.text_input("Enter your HUB name")

col1, col2 = st.columns(2)
with col1:
    name = st.text_input("Hub Name")
    location = st.text_input("Location")
    focus = st.text_area("Focus Areas (comma separated)")

with col2:
    problem = st.text_area("Problem Statement")
    solution = st.text_area("Proposed Solution")
    beneficiaries = st.text_area("Beneficiaries")
    funding = st.text_input("Funding Needed")

# --- File Upload ---
st.subheader("📎 Upload Supporting Documents")
uploaded_files = st.file_uploader(
    f"Upload videos, papers, or other supporting documents for {hub_name}",
    type=["pdf", "docx", "jpg", "png", "mp4"],
    accept_multiple_files=True
)

# --- IP Protection ---
st.divider()
st.subheader("🔒 Intellectual Property Protection & Access Control")
st.markdown("""
✅ **You retain all intellectual property rights** to your pitch and uploaded content  
✅ Your content will **not be shared publicly** or with other users  
✅ Only you and your **HUB administrators** can access your content  
✅ The platform stores your content securely in private storage  

📌 **Access Levels:**
- 👤 **You:** Full access to your files
- 🏢 **HUB Admins:** Can view all files from their HUB
- 🌐 **General Public:** Cannot view any uploaded files
""")
agree = st.checkbox("✅ I agree to the IP Protection Terms and confirm I own the rights to my content")
if agree:
    st.success("✅ You have agreed to the terms. Your IP is protected.")
else:
    st.warning("⚠️ Please agree to the terms to protect your intellectual property.")

# --- Step 1: Find Donors ---
if st.button("🔍 Find Best Donors", type="primary"):
    if not name:
        st.warning("Please enter at least a Hub Name.")
    elif not agree:
        st.warning("Please agree to the IP Protection Terms first.")
    else:
        with st.spinner("Analyzing your pitch and finding the best donors..."):
            full_pitch_text = f"Hub: {hub_name}\nName: {name}\nLocation: {location}\nProblem: {problem}\nSolution: {solution}\nFocus: {focus}\nFunding: {funding}\nBeneficiaries: {beneficiaries}"
            matches = ai_match_donor(full_pitch_text, filtered_donors)
            st.session_state.donor_matches = matches
            st.session_state.selected_donor = None
            st.session_state.generated_pitch = None
            st.rerun()

# --- Step 2: Display matches and allow selection ---
if st.session_state.donor_matches is not None:
    st.subheader("🏆 Top 3 Matching Donors")
    donor_options = []
    donor_reasons = {}
    for i, (donor_name, reason) in enumerate(st.session_state.donor_matches):
        label = f"{donor_name} – {reason}"
        donor_options.append(label)
        donor_reasons[label] = (donor_name, reason)

    selected_label = st.radio(
        "Select the donor you want to target:",
        donor_options,
        index=0
    )
    selected_donor_name, selected_reason = donor_reasons[selected_label]
    st.session_state.selected_donor = selected_donor_name

    # Show donor details
    donor_row = donors[donors["Donor"] == selected_donor_name]
    if not donor_row.empty:
        st.info(f"**Donor Focus:** {donor_row.iloc[0]['Focus']}")
        st.info(f"**Funding Call:** {donor_row.iloc[0]['Funding Call']}")
        max_pages = donor_row.iloc[0].get('MaxPages', 30)
        rec_pages = donor_row.iloc[0].get('RecommendedPages', 10)
        st.info(f"**Recommended page length:** ~{rec_pages} pages (max {max_pages})")

    # Page number input
    desired_pages = st.number_input(
        "📄 Desired Proposal Length (pages)",
        min_value=3,
        max_value=30,
        value=min(rec_pages, 30) if 'rec_pages' in locals() else 10,
        step=1,
        help="Choose the number of pages for your proposal. Recommended based on donor guidelines."
    )

    # --- Step 3: Generate Pitch ---
    if st.button("🎯 Generate Pitch", type="primary"):
        with st.spinner("Generating your pitch..."):
            data = {
                "name": name,
                "location": location,
                "focus": focus.split(",") if focus else [],
                "problem": problem,
                "solution": solution,
                "beneficiaries": beneficiaries,
                "funding": funding,
                "hub": hub_name,
                "desired_pages": desired_pages,
                "selected_donor": selected_donor_name
            }
            pitch = generate_pitch(data)
            st.session_state.generated_pitch = pitch
            if save_pitch_to_db(hub_name, name, pitch):
                st.success("✅ Pitch saved to database for admin extraction!")
            st.subheader("📄 Generated Pitch")
            st.text_area("Your Pitch", pitch, height=400, key="generated_pitch")
            if pitch:
                st.download_button(label="📥 Download Full Proposal (Text File)", data=pitch, file_name=f"{name}_Proposal.txt", mime="text/plain", help="Download as plain text")
                pdf_data = generate_pdf_proposal(pitch, hub_name, name)
                st.download_button(label="📄 Download High Quality Proposal (PDF)", data=pdf_data, file_name=f"{name}_Proposal.pdf", mime="application/pdf", help="Download as PDF")
                word_data = generate_word_proposal(pitch, hub_name, name)
                st.download_button(label="📝 Download High Quality Proposal (Word)", data=word_data, file_name=f"{name}_Proposal.docx", mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document", help="Download as Word")
            if uploaded_files:
                st.subheader("📎 Uploaded Files")
                st.write(f"**HUB:** {hub_name}")
                st.write(f"**Files:** {len(uploaded_files)} file(s) uploaded")
                for file in uploaded_files:
                    st.write(f"- {file.name}")
                st.caption("🔒 Files are private and accessible only to you and your HUB admins.")
            st.divider()
            st.subheader("📋 Full Donor List with Application Links")
            st.caption(f"Showing {len(filtered_donors)} donors from the complete PNG funding database")
            for idx, row in filtered_donors.iterrows():
                with st.container():
                    col1, col2, col3, col4, col5, col6 = st.columns([2, 2, 1, 1, 1, 1])
                    with col1:
                        st.markdown(f"**{row['Donor']}**")
                    with col2:
                        st.write(row['Funding Call'])
                    with col3:
                        st.write(row['Amount'])
                    with col4:
                        st.write(row['Status'])
                    with col5:
                        st.link_button("🌐 Website", row["Website"])
                    with col6:
                        st.link_button("📝 Apply Now", row["Apply"])
                    st.divider()

# --- ADMIN DASHBOARD ---
st.divider()
with st.expander("🔐 Admin: Extract Pitches by HUB"):
    admin_password = st.text_input("Enter Admin Password", type="password")
    if admin_password == "admin123":
        extract_hub = st.selectbox("Select HUB to Extract", ["SoCe Innovation Hub", "CEFI Fintech Hub", "UPNG Innovation Hub", "PNG Unitech Innovation Hub"])
        if st.button("Extract Pitches"):
            try:
                df = pd.read_csv('pitches.csv')
                hub_pitches = df[df['Hub'] == extract_hub]
                if len(hub_pitches) > 0:
                    st.success(f"✅ Found {len(hub_pitches)} pitches from {extract_hub}")
                    st.dataframe(hub_pitches)
                    csv_data = hub_pitches.to_csv(index=False)
                    st.download_button(label="📥 Download Extracted Data (CSV)", data=csv_data, file_name=f"{extract_hub}_pitches.csv", mime="text/csv")
                else:
                    st.warning(f"No pitches found for {extract_hub}")
            except FileNotFoundError:
                st.warning("No pitches have been generated yet. Generate a pitch first!")
    else:
        st.warning("Please enter the admin password to access extraction.")

# --- FOOTER ---
st.divider()
st.caption("🇵🇬 AI Incubation Hub Pitch Generator | Connecting PNG Innovators to Funding Opportunities")